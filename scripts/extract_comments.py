#!/usr/bin/env python3
"""
extract_comments.py — Robust Word comment extractor with anchor resolution.

WHY THIS EXISTS
---------------
The external `office-word-mcp-server` tool `get_all_comments` extracts each
comment's *text* but never resolves WHERE the comment is anchored: it always
returns `paragraph_index: null` and `reference_text: ""`. That makes it
impossible to act on a comment ("clarify the sampling strategy") because the
agent cannot tell which sentence in the body the reviewer is pointing at. Its
`get_comments_for_paragraph` is consequently always empty.

This script replaces that read path. For every comment it resolves:
  - text            : the comment content
  - author / initials / date
  - reference_text  : the exact body text the comment is anchored to
                      (the run text between <w:commentRangeStart> and
                       <w:commentRangeEnd> for the comment's id)
  - paragraph_index : body-order index of the paragraph where the anchor
                      starts (matches python-docx doc.paragraphs ordering;
                      table-cell paragraphs are flagged via `in_table`)
  - resolved        : True if the thread is marked done (commentsExtended)
  - parent_id       : the comment this one replies to (None for top-level)

Usage:
    python extract_comments.py <file.docx> [--author NAME] [--paragraph N]
    python extract_comments.py <file.docx> --markdown

Output: JSON to stdout (or a Markdown table with --markdown).
Exit code 0 on success, 1 on error (error JSON still printed to stdout).
"""
import sys
import os
import json
import argparse

from docx import Document
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"


def _q(ns, tag):
    return f"{{{ns}}}{tag}"


def _find_part(doc, suffix):
    """Find a related part whose reltype's last segment == suffix (e.g. 'comments')."""
    for rel in doc.part.rels.values():
        if rel.reltype.split('/')[-1] == suffix:
            try:
                return rel.target_part
            except Exception:
                return None
    return None


def _part_root(part):
    """Return a plain lxml root element for a part by parsing its serialized
    bytes. We go through `.blob` (available on every part) rather than
    `.element` so the result is a vanilla lxml element whose `.xpath()` accepts
    a `namespaces=` kwarg — python-docx's own element xpath does not. This also
    transparently handles generic parts like commentsExtended that python-docx
    does not model with an `.element`."""
    if part is None:
        return None
    try:
        return etree.fromstring(part.blob)
    except Exception:
        return None


def _comment_meta(doc):
    """id -> {comment_id, author, initials, date, text} from word/comments.xml."""
    meta = {}
    root = _part_root(_find_part(doc, 'comments'))
    if root is None:
        return meta, {}
    paraid_to_cid = {}  # w14:paraId of comment's first paragraph -> comment id
    NSMAP = {'w': W, 'w14': W14}
    for ce in root.xpath('.//w:comment', namespaces=NSMAP):
        cid = ce.get(_q(W, 'id'))
        if cid is None:
            continue
        texts = ce.xpath('.//w:t', namespaces=NSMAP)
        meta[cid] = {
            'comment_id': cid,
            'author': ce.get(_q(W, 'author'), 'Unknown'),
            'initials': ce.get(_q(W, 'initials'), ''),
            'date': ce.get(_q(W, 'date'), '') or None,
            'text': ''.join(t.text or '' for t in texts).strip(),
            # parentId is present in some producers directly on the comment
            'parent_id': ce.get(_q(W, 'parentId')),
            'resolved': False,
        }
        first_p = ce.find(_q(W, 'p'))
        if first_p is not None:
            pid = first_p.get(_q(W14, 'paraId'))
            if pid:
                paraid_to_cid[pid] = cid
    return meta, paraid_to_cid


def _apply_extended(doc, meta, paraid_to_cid):
    """Merge resolved status + threading from word/commentsExtended.xml."""
    root = _part_root(_find_part(doc, 'commentsExtended'))
    if root is None:
        return
    for ex in root.xpath('.//w15:commentEx', namespaces={'w15': W15}):
        pid = ex.get(_q(W15, 'paraId'))
        cid = paraid_to_cid.get(pid)
        if cid is None or cid not in meta:
            continue
        done = ex.get(_q(W15, 'done'))
        if done in ('1', 'true', 'True'):
            meta[cid]['resolved'] = True
        parent_pid = ex.get(_q(W15, 'paraIdParent'))
        if parent_pid and meta[cid].get('parent_id') is None:
            meta[cid]['parent_id'] = paraid_to_cid.get(parent_pid)


def _resolve_anchors(doc):
    """
    Walk the body in document order, tracking open comment ranges, to collect
    the anchored text and starting paragraph index for each comment id.

    paragraph_index counts <w:p> elements in body order. To stay consistent
    with python-docx's doc.paragraphs (which excludes table-cell paragraphs),
    table paragraphs are tracked separately and reported with in_table=True and
    a paragraph_index of None.
    """
    anchors = {}            # cid -> {paragraph_index, reference_text, in_table}
    open_ranges = {}        # cid -> {'frags': [...], 'pidx': int|None, 'in_table': bool}
    body = doc.element.body
    body_para_idx = -1      # mirrors doc.paragraphs indexing (top-level only)
    table_depth = 0

    for el in body.iter():
        tag = el.tag
        if tag == _q(W, 'tbl'):
            # Note: nested table entry; we still count via tc depth below.
            pass
        if tag == _q(W, 'p'):
            # Determine if this paragraph is inside a table cell.
            in_table = _is_in_table(el)
            if not in_table:
                body_para_idx += 1
            # Stamp current paragraph context for any open range that starts here.
            el_pidx = None if in_table else body_para_idx
            el_in_table = in_table
            # Make available to range starts encountered within this paragraph.
            _CTX['pidx'] = el_pidx
            _CTX['in_table'] = el_in_table
        elif tag == _q(W, 'commentRangeStart'):
            cid = el.get(_q(W, 'id'))
            if cid is not None:
                open_ranges[cid] = {
                    'frags': [],
                    'pidx': _CTX.get('pidx'),
                    'in_table': _CTX.get('in_table', False),
                }
        elif tag == _q(W, 'commentRangeEnd'):
            cid = el.get(_q(W, 'id'))
            if cid in open_ranges:
                r = open_ranges.pop(cid)
                anchors[cid] = {
                    'paragraph_index': r['pidx'],
                    'reference_text': ''.join(r['frags']).strip(),
                    'in_table': r['in_table'],
                }
        elif tag == _q(W, 't'):
            if open_ranges:
                txt = el.text or ''
                for r in open_ranges.values():
                    r['frags'].append(txt)

    # Any unclosed ranges (malformed doc) — record what we have.
    for cid, r in open_ranges.items():
        anchors[cid] = {
            'paragraph_index': r['pidx'],
            'reference_text': ''.join(r['frags']).strip(),
            'in_table': r['in_table'],
        }
    return anchors


# Shared paragraph context for range-start stamping during a single iter() walk.
_CTX = {}


def _is_in_table(p_el):
    """True if a <w:p> element has a <w:tc> ancestor."""
    parent = p_el.getparent()
    while parent is not None:
        if parent.tag == _q(W, 'tc'):
            return True
        parent = parent.getparent()
    return False


def extract_comments(path):
    """Return a list of comment dicts with resolved anchors, or raise."""
    doc = Document(path)
    meta, paraid_to_cid = _comment_meta(doc)
    if not meta:
        return []
    _apply_extended(doc, meta, paraid_to_cid)
    global _CTX
    _CTX = {}
    anchors = _resolve_anchors(doc)

    out = []
    for cid, m in meta.items():
        a = anchors.get(cid, {
            'paragraph_index': None,
            'reference_text': '',
            'in_table': False,
        })
        rec = dict(m)
        rec['paragraph_index'] = a['paragraph_index']
        rec['reference_text'] = a['reference_text']
        rec['in_table'] = a['in_table']
        rec['is_reply'] = rec.get('parent_id') is not None
        out.append(rec)

    # Stable order: by starting paragraph, then by comment id.
    def _sortkey(r):
        pi = r['paragraph_index']
        return (0 if pi is not None else 1, pi if pi is not None else 0,
                int(r['comment_id']) if str(r['comment_id']).isdigit() else 0)
    out.sort(key=_sortkey)
    return out


def _to_markdown(comments):
    rows = ["| # | Author | Resolved | Anchored Text | Comment | Location |",
            "|---|--------|----------|---------------|---------|----------|"]
    for i, c in enumerate(comments, 1):
        loc = ("table" if c['in_table']
               else (f"¶{c['paragraph_index']}" if c['paragraph_index'] is not None
                     else "?"))
        anchor = (c['reference_text'][:40] + '…') if len(c['reference_text']) > 40 else c['reference_text']
        text = (c['text'][:50] + '…') if len(c['text']) > 50 else c['text']
        reply = " (reply)" if c['is_reply'] else ""
        rows.append(f"| {i} | {c['author']}{reply} | {'✓' if c['resolved'] else '—'} "
                    f"| {anchor or '—'} | {text} | {loc} |")
    return "\n".join(rows)


def main(argv=None):
    ap = argparse.ArgumentParser(description="Extract Word comments with anchor resolution.")
    ap.add_argument('file', help='Path to the .docx file')
    ap.add_argument('--author', help='Filter to comments by this author (case-insensitive)')
    ap.add_argument('--paragraph', type=int, help='Filter to comments anchored at this paragraph index')
    ap.add_argument('--markdown', action='store_true', help='Emit a Markdown table instead of JSON')
    args = ap.parse_args(argv)

    path = args.file
    if not path.lower().endswith('.docx'):
        path = path + '.docx'
    if not os.path.exists(path):
        print(json.dumps({'success': False, 'error': f'File not found: {path}'}, ensure_ascii=False, indent=2))
        return 1

    try:
        comments = extract_comments(path)
    except Exception as e:
        print(json.dumps({'success': False, 'error': f'Failed to extract comments: {e}'},
                         ensure_ascii=False, indent=2))
        return 1

    if args.author:
        a = args.author.lower()
        comments = [c for c in comments if c['author'].lower() == a]
    if args.paragraph is not None:
        comments = [c for c in comments if c['paragraph_index'] == args.paragraph]

    if args.markdown:
        print(_to_markdown(comments) if comments else "_(no comments found)_")
    else:
        print(json.dumps({
            'success': True,
            'total_comments': len(comments),
            'comments': comments,
        }, ensure_ascii=False, indent=2))
    return 0


if __name__ == '__main__':
    sys.exit(main())
