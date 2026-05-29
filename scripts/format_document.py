#!/usr/bin/env python3
"""
format_document.py — Academic paper formatting engine for Word documents.

Handles paragraph-level and page-level formatting that MCP tools cannot:
  page-setup      Page size, margins, orientation
  styles          Modify built-in style definitions (font, spacing, indent)
  paragraph       Apply paragraph formatting by style name or index range
  header-footer   Header/footer content and page numbers
  section         Section breaks and per-section formatting
  toc             Insert Table of Contents field
  apply-spec      Apply all formatting from a JSON spec file

Safety: creates backup before any modification, atomic write, file lock detection.
"""

import argparse
import json
import os
import platform
import shutil
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor

# ─── Chinese font size mapping ───────────────────────────────────────────────

CN_FONT_SIZES = {
    "初号": 42, "小初": 36,
    "一号": 26, "小一": 24,
    "二号": 22, "小二": 18,
    "三号": 16, "小三": 15,
    "四号": 14, "小四": 12,
    "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5,
    "七号": 5.5, "八号": 5,
}

PAPER_SIZES = {
    "A4": (Cm(21.0), Cm(29.7)),
    "A3": (Cm(29.7), Cm(42.0)),
    "B5": (Cm(17.6), Cm(25.0)),
    "Letter": (Inches(8.5), Inches(11)),
    "Legal": (Inches(8.5), Inches(14)),
    "16K": (Cm(18.4), Cm(26.0)),
}

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}

CN_FONT_PAIRS = {
    "宋体": "Times New Roman", "SimSun": "Times New Roman",
    "黑体": "Arial", "SimHei": "Arial",
    "楷体": "Times New Roman", "KaiTi": "Times New Roman",
    "仿宋": "Times New Roman", "FangSong": "Times New Roman",
    "微软雅黑": "Arial", "Microsoft YaHei": "Arial",
    "华文中宋": "Times New Roman", "STZhongsong": "Times New Roman",
    "华文楷体": "Times New Roman", "STKaiti": "Times New Roman",
    "华文宋体": "Times New Roman", "STSong": "Times New Roman",
    "华文仿宋": "Times New Roman", "STFangsong": "Times New Roman",
    "华文细黑": "Arial", "STXihei": "Arial",
    "方正小标宋": "Times New Roman", "方正仿宋": "Times New Roman",
    "方正楷体": "Times New Roman", "方正黑体": "Arial",
}


# ─── Utility functions ───────────────────────────────────────────────────────

def parse_size(val):
    """Parse size string to docx.shared unit.
    Supports: '12pt', '2.54cm', '1in', '25.4mm', '小四', '2char', raw number (pt).
    """
    if val is None:
        return None
    val = str(val).strip()
    if val in CN_FONT_SIZES:
        return Pt(CN_FONT_SIZES[val])
    if val.endswith("pt"):
        return Pt(float(val[:-2]))
    if val.endswith("cm"):
        return Cm(float(val[:-2]))
    if val.endswith("mm"):
        return Mm(float(val[:-2]))
    if val.endswith("in"):
        return Inches(float(val[:-2]))
    try:
        return Pt(float(val))
    except ValueError:
        raise ValueError(f"Cannot parse size: '{val}'. Use 12pt, 2.54cm, 1in, 小四, etc.")


def parse_margins(margin_str):
    """Parse margin string 'top,bottom,left,right' into tuple of sizes."""
    parts = [s.strip() for s in margin_str.split(",")]
    if len(parts) == 1:
        m = parse_size(parts[0])
        return m, m, m, m
    if len(parts) == 2:
        tb, lr = parse_size(parts[0]), parse_size(parts[1])
        return tb, tb, lr, lr
    if len(parts) == 4:
        return tuple(parse_size(p) for p in parts)
    raise ValueError(f"Margins must be 1, 2, or 4 values: '{margin_str}'")


def check_file_lock(filepath):
    """Check if a Word lock file exists (indicating the file is open in Word)."""
    p = Path(filepath)
    lock_file = p.parent / f"~${p.name}"
    if lock_file.exists():
        print(f"ERROR: File is locked by Word (lock file: {lock_file})", file=sys.stderr)
        print("Close the document in Word before running this script.", file=sys.stderr)
        sys.exit(2)


def create_backup(filepath):
    """Create a timestamped backup of the document."""
    p = Path(filepath)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = p.parent / f"{p.stem}_backup_{ts}{p.suffix}"
    shutil.copy2(filepath, backup)
    print(f"Backup: {backup}")
    return backup


def verify_docx(filepath):
    """Verify the output file is a valid zip (docx) archive."""
    try:
        with zipfile.ZipFile(filepath, 'r') as zf:
            if "word/document.xml" not in zf.namelist():
                raise ValueError("Missing word/document.xml")
        return True
    except Exception as e:
        print(f"ERROR: Output file verification failed: {e}", file=sys.stderr)
        return False


def atomic_save(doc, filepath):
    """Save document atomically: write to temp file, then rename."""
    p = Path(filepath)
    fd, tmp = tempfile.mkstemp(suffix=".docx", dir=p.parent)
    os.close(fd)
    try:
        doc.save(tmp)
        if not verify_docx(tmp):
            os.unlink(tmp)
            print("ERROR: Save aborted — output file is corrupted.", file=sys.stderr)
            sys.exit(1)
        shutil.move(tmp, filepath)
    except Exception:
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise


def set_rfonts(element, cn_font=None, en_font=None):
    """Set w:rFonts on an XML element, handling both ascii and eastAsia."""
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        element.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    if en_font:
        rFonts.set(qn('w:ascii'), en_font)
        rFonts.set(qn('w:hAnsi'), en_font)
        rFonts.set(qn('w:cs'), en_font)
    if cn_font:
        rFonts.set(qn('w:eastAsia'), cn_font)

    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme']:
        if rFonts.get(qn(attr)):
            del rFonts.attrib[qn(attr)]


# ─── Core formatting engine ─────────────────────────────────────────────────

class FormatEngine:
    def __init__(self, filepath, backup=True):
        self.filepath = str(filepath)
        check_file_lock(self.filepath)
        if backup:
            self.backup_path = create_backup(self.filepath)
        self.doc = Document(self.filepath)
        self.changes = []

    def _log(self, action):
        self.changes.append(action)
        print(f"  ✓ {action}")

    def save(self):
        atomic_save(self.doc, self.filepath)
        print(f"\nSaved: {self.filepath}")
        print(f"Changes applied: {len(self.changes)}")

    # ── page-setup ────────────────────────────────────────────────────────

    def page_setup(self, size=None, margins=None, orientation=None):
        """Set page size, margins, and orientation for all sections."""
        print("\n[page-setup]")
        for i, section in enumerate(self.doc.sections):
            label = f"Section {i+1}"

            if size:
                size_upper = size.upper()
                if size_upper in PAPER_SIZES:
                    w, h = PAPER_SIZES[size_upper]
                else:
                    raise ValueError(f"Unknown paper size: {size}. Use: {', '.join(PAPER_SIZES)}")
                section.page_width = w
                section.page_height = h
                self._log(f"{label}: page size → {size_upper} ({w}, {h})")

            if orientation:
                orient = orientation.lower()
                if orient == "portrait":
                    section.orientation = WD_ORIENT.PORTRAIT
                    if section.page_width > section.page_height:
                        section.page_width, section.page_height = section.page_height, section.page_width
                elif orient == "landscape":
                    section.orientation = WD_ORIENT.LANDSCAPE
                    if section.page_height > section.page_width:
                        section.page_width, section.page_height = section.page_height, section.page_width
                else:
                    raise ValueError(f"Orientation must be 'portrait' or 'landscape', got: {orient}")
                self._log(f"{label}: orientation → {orient}")

            if margins:
                top, bottom, left, right = parse_margins(margins)
                section.top_margin = top
                section.bottom_margin = bottom
                section.left_margin = left
                section.right_margin = right
                self._log(f"{label}: margins → T={top} B={bottom} L={left} R={right}")

    # ── styles ────────────────────────────────────────────────────────────

    def modify_style(self, style_name, cn_font=None, en_font=None, font_size=None,
                     bold=None, italic=None, color=None, alignment=None,
                     line_spacing=None, line_spacing_rule=None,
                     space_before=None, space_after=None,
                     first_indent=None, hanging_indent=None,
                     keep_with_next=None):
        """Modify an existing style definition. Changes propagate to all paragraphs using this style."""
        print(f"\n[styles] Modifying: '{style_name}'")

        try:
            style = self.doc.styles[style_name]
        except KeyError:
            print(f"  WARNING: Style '{style_name}' not found. Creating it.")
            from docx.enum.style import WD_STYLE_TYPE
            style = self.doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

        # Font properties
        if en_font:
            style.font.name = en_font
            self._log(f"font.ascii/hAnsi → {en_font}")
        if cn_font:
            set_rfonts(style.element, cn_font=cn_font, en_font=en_font)
            self._log(f"font.eastAsia → {cn_font}")
        elif en_font:
            auto_cn = CN_FONT_PAIRS.get(en_font)
            if auto_cn:
                set_rfonts(style.element, cn_font=auto_cn, en_font=en_font)

        if font_size:
            style.font.size = parse_size(font_size)
            self._log(f"font.size → {font_size}")
        if bold is not None:
            style.font.bold = bold
            self._log(f"font.bold → {bold}")
        if italic is not None:
            style.font.italic = italic
            self._log(f"font.italic → {italic}")
        if color:
            style.font.color.rgb = RGBColor.from_string(color)
            self._log(f"font.color → #{color}")

        # Paragraph properties
        pf = style.paragraph_format

        if alignment and alignment.lower() in ALIGNMENT_MAP:
            pf.alignment = ALIGNMENT_MAP[alignment.lower()]
            self._log(f"alignment → {alignment}")

        if line_spacing is not None:
            ls = float(line_spacing)
            if line_spacing_rule == "exact":
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = parse_size(str(line_spacing))
                self._log(f"line_spacing → exactly {line_spacing}")
            elif line_spacing_rule == "at_least":
                pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
                pf.line_spacing = parse_size(str(line_spacing))
                self._log(f"line_spacing → at least {line_spacing}")
            else:
                if ls <= 0:
                    raise ValueError("Line spacing multiplier must be > 0")
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = ls
                self._log(f"line_spacing → {ls}x")

        if space_before is not None:
            pf.space_before = parse_size(space_before)
            self._log(f"space_before → {space_before}")
        if space_after is not None:
            pf.space_after = parse_size(space_after)
            self._log(f"space_after → {space_after}")

        if first_indent is not None:
            fi = str(first_indent)
            if fi.endswith("char"):
                chars = int(float(fi.replace("char", "")) * 100)
                pPr = style.element.get_or_add_pPr()
                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = OxmlElement('w:ind')
                    pPr.append(ind)
                ind.set(qn('w:firstLineChars'), str(chars))
                if ind.get(qn('w:firstLine')):
                    del ind.attrib[qn('w:firstLine')]
                self._log(f"first_indent → {fi}")
            else:
                pf.first_line_indent = parse_size(fi)
                self._log(f"first_indent → {fi}")

        if hanging_indent is not None:
            pf.first_line_indent = -parse_size(hanging_indent)
            self._log(f"hanging_indent → {hanging_indent}")

        if keep_with_next is not None:
            pf.keep_with_next = keep_with_next
            self._log(f"keep_with_next → {keep_with_next}")

    # ── paragraph ─────────────────────────────────────────────────────────

    def paragraph_format_batch(self, style_filter=None, index_range=None,
                               alignment=None, line_spacing=None,
                               space_before=None, space_after=None,
                               first_indent=None):
        """Apply paragraph formatting to paragraphs matching a filter."""
        print(f"\n[paragraph] Filter: style={style_filter}, range={index_range}")

        paragraphs = list(self.doc.paragraphs)
        targets = []

        for i, p in enumerate(paragraphs):
            if index_range:
                start, end = index_range
                if not (start <= i <= end):
                    continue
            if style_filter:
                if p.style and p.style.name != style_filter:
                    continue
            targets.append((i, p))

        if not targets:
            print("  WARNING: No paragraphs matched the filter.")
            return

        count = 0
        for i, p in targets:
            pf = p.paragraph_format

            if alignment and alignment.lower() in ALIGNMENT_MAP:
                pf.alignment = ALIGNMENT_MAP[alignment.lower()]
            if line_spacing is not None:
                ls = float(line_spacing)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = ls
            if space_before is not None:
                pf.space_before = parse_size(space_before)
            if space_after is not None:
                pf.space_after = parse_size(space_after)
            if first_indent is not None:
                fi = str(first_indent)
                if fi.endswith("char"):
                    chars = int(float(fi.replace("char", "")) * 100)
                    pPr = p._element.get_or_add_pPr()
                    ind = pPr.find(qn('w:ind'))
                    if ind is None:
                        ind = OxmlElement('w:ind')
                        pPr.append(ind)
                    ind.set(qn('w:firstLineChars'), str(chars))
                    if ind.get(qn('w:firstLine')):
                        del ind.attrib[qn('w:firstLine')]
                else:
                    pf.first_line_indent = parse_size(fi)
            count += 1

        self._log(f"Applied formatting to {count} paragraphs")

    # ── header-footer ─────────────────────────────────────────────────────

    def header_footer(self, header_text=None, footer_text=None,
                      header_font=None, header_size=None,
                      footer_font=None, footer_size=None,
                      page_number=False, page_number_pos="center",
                      different_first=False):
        """Set header/footer content and page numbers."""
        print("\n[header-footer]")

        for i, section in enumerate(self.doc.sections):
            label = f"Section {i+1}"

            if different_first:
                section.different_first_page_header_footer = True
                self._log(f"{label}: different first page → enabled")

            # Header
            if header_text is not None:
                header = section.header
                header.is_linked_to_previous = False
                if header.paragraphs:
                    hp = header.paragraphs[0]
                    hp.clear()
                else:
                    hp = header.add_paragraph()

                run = hp.add_run(header_text)
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if header_font:
                    cn, en = self._split_font(header_font)
                    run.font.name = en
                    set_rfonts(run._element, cn_font=cn, en_font=en)
                if header_size:
                    run.font.size = parse_size(header_size)

                self._log(f"{label}: header → '{header_text}'")

            # Footer
            if footer_text is not None or page_number:
                footer = section.footer
                footer.is_linked_to_previous = False
                if footer.paragraphs:
                    fp = footer.paragraphs[0]
                    fp.clear()
                else:
                    fp = footer.add_paragraph()

                if footer_text:
                    run = fp.add_run(footer_text)
                    if footer_font:
                        cn, en = self._split_font(footer_font)
                        run.font.name = en
                        set_rfonts(run._element, cn_font=cn, en_font=en)
                    if footer_size:
                        run.font.size = parse_size(footer_size)

                if page_number:
                    self._insert_page_number(fp, footer_font, footer_size)

                align = page_number_pos.lower() if page_number_pos else "center"
                if align in ALIGNMENT_MAP:
                    fp.alignment = ALIGNMENT_MAP[align]

                self._log(f"{label}: footer → text='{footer_text or ''}', page_number={page_number}")

    def _split_font(self, font_str):
        """Split 'cn_font,en_font' or return (font, auto_pair)."""
        if "," in font_str:
            parts = [f.strip() for f in font_str.split(",")]
            return parts[0], parts[1]
        f = font_str.strip()
        pair = CN_FONT_PAIRS.get(f, f)
        return f, pair

    def _insert_page_number(self, paragraph, font_str=None, size_str=None):
        """Insert a PAGE field code into a paragraph."""
        run = paragraph.add_run()
        if font_str:
            cn, en = self._split_font(font_str)
            run.font.name = en
            set_rfonts(run._element, cn_font=cn, en_font=en)
        if size_str:
            run.font.size = parse_size(size_str)

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar_begin)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._element.append(instrText)

        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar_end)

    # ── section ───────────────────────────────────────────────────────────

    def add_section_break(self, after_paragraph=None, after_text=None,
                          page_number_start=None, orientation=None):
        """Add a section break and configure the new section."""
        print("\n[section]")

        target_idx = None
        if after_paragraph is not None:
            target_idx = int(after_paragraph)
        elif after_text:
            for i, p in enumerate(self.doc.paragraphs):
                if after_text in p.text:
                    target_idx = i
                    break
            if target_idx is None:
                print(f"  WARNING: Text '{after_text}' not found.")
                return

        if target_idx is not None:
            p = self.doc.paragraphs[target_idx]
            pPr = p._element.get_or_add_pPr()
            sectPr = OxmlElement('w:sectPr')
            sect_type = OxmlElement('w:type')
            sect_type.set(qn('w:val'), 'nextPage')
            sectPr.append(sect_type)
            pPr.append(sectPr)
            self._log(f"Section break after paragraph {target_idx}")

        if page_number_start is not None:
            sections = list(self.doc.sections)
            last_section = sections[-1]
            pgNumType = last_section._sectPr.find(qn('w:pgNumType'))
            if pgNumType is None:
                pgNumType = OxmlElement('w:pgNumType')
                last_section._sectPr.append(pgNumType)
            pgNumType.set(qn('w:start'), str(page_number_start))
            self._log(f"Page number restart at {page_number_start}")

    # ── toc ───────────────────────────────────────────────────────────────

    def insert_toc(self, position=0, levels=3, title="目录"):
        """Insert a Table of Contents field at the given paragraph position."""
        print("\n[toc]")

        paragraphs = self.doc.paragraphs
        if isinstance(position, str):
            for i, p in enumerate(paragraphs):
                if position in p.text:
                    position = i
                    break

        if isinstance(position, str):
            print(f"  WARNING: Position text '{position}' not found. Inserting at beginning.")
            position = 0

        ref_paragraph = paragraphs[min(position, len(paragraphs) - 1)]
        ref_element = ref_paragraph._element

        # TOC title paragraph
        toc_title = OxmlElement('w:p')
        toc_title_pPr = OxmlElement('w:pPr')
        toc_title_style = OxmlElement('w:pStyle')
        toc_title_style.set(qn('w:val'), 'TOCHeading')
        toc_title_pPr.append(toc_title_style)
        toc_title.append(toc_title_pPr)
        toc_title_run = OxmlElement('w:r')
        toc_title_text = OxmlElement('w:t')
        toc_title_text.text = title
        toc_title_run.append(toc_title_text)
        toc_title.append(toc_title_run)
        ref_element.addprevious(toc_title)

        # SDT (Structured Document Tag) for TOC
        sdt = OxmlElement('w:sdt')
        sdtPr = OxmlElement('w:sdtPr')
        docPartObj = OxmlElement('w:docPartObj')
        docPartGallery = OxmlElement('w:docPartGallery')
        docPartGallery.set(qn('w:val'), 'Table of Contents')
        docPartObj.append(docPartGallery)
        sdtPr.append(docPartObj)
        sdt.append(sdtPr)

        sdtContent = OxmlElement('w:sdtContent')
        toc_p = OxmlElement('w:p')
        toc_r = OxmlElement('w:r')

        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        toc_r.append(fldChar_begin)
        toc_p.append(toc_r)

        toc_r2 = OxmlElement('w:r')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' TOC \\o "1-{levels}" \\h \\z \\u '
        toc_r2.append(instrText)
        toc_p.append(toc_r2)

        toc_r3 = OxmlElement('w:r')
        fldChar_separate = OxmlElement('w:fldChar')
        fldChar_separate.set(qn('w:fldCharType'), 'separate')
        toc_r3.append(fldChar_separate)
        toc_p.append(toc_r3)

        toc_r4 = OxmlElement('w:r')
        toc_text = OxmlElement('w:t')
        toc_text.text = "[请在 Word 中右键点击此处 → 更新域 以生成目录]"
        toc_r4.append(toc_text)
        toc_p.append(toc_r4)

        toc_r5 = OxmlElement('w:r')
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        toc_r5.append(fldChar_end)
        toc_p.append(toc_r5)

        sdtContent.append(toc_p)
        sdt.append(sdtContent)

        ref_element.addprevious(sdt)
        self._log(f"TOC inserted at paragraph {position}, levels=1-{levels}")
        self._log("NOTE: Open in Word → right-click TOC → Update Field to populate")

    # ── apply-spec ────────────────────────────────────────────────────────

    def apply_spec(self, spec_path):
        """Apply all formatting from a JSON spec file."""
        print(f"\n[apply-spec] Reading: {spec_path}")
        with open(spec_path, 'r', encoding='utf-8') as f:
            spec = json.load(f)

        # Page setup
        page = spec.get("page", {})
        if page:
            self.page_setup(
                size=page.get("size"),
                margins=page.get("margins"),
                orientation=page.get("orientation"),
            )

        # Style modifications
        styles = spec.get("styles", {})
        for style_name, props in styles.items():
            self.modify_style(
                style_name,
                cn_font=props.get("cn_font"),
                en_font=props.get("en_font"),
                font_size=props.get("font_size"),
                bold=props.get("bold"),
                italic=props.get("italic"),
                color=props.get("color"),
                alignment=props.get("alignment"),
                line_spacing=props.get("line_spacing"),
                line_spacing_rule=props.get("line_spacing_rule"),
                space_before=props.get("space_before"),
                space_after=props.get("space_after"),
                first_indent=props.get("first_indent"),
                keep_with_next=props.get("keep_with_next"),
            )

        # Paragraph overrides (for specific ranges or patterns)
        paragraphs = spec.get("paragraphs", [])
        for pspec in paragraphs:
            idx_range = None
            if "range" in pspec:
                r = pspec["range"]
                idx_range = (r[0], r[1])
            self.paragraph_format_batch(
                style_filter=pspec.get("style"),
                index_range=idx_range,
                alignment=pspec.get("alignment"),
                line_spacing=pspec.get("line_spacing"),
                space_before=pspec.get("space_before"),
                space_after=pspec.get("space_after"),
                first_indent=pspec.get("first_indent"),
            )

        # Header/footer
        hf = spec.get("header_footer", {})
        if hf:
            self.header_footer(
                header_text=hf.get("header_text"),
                footer_text=hf.get("footer_text"),
                header_font=hf.get("header_font"),
                header_size=hf.get("header_size"),
                footer_font=hf.get("footer_font"),
                footer_size=hf.get("footer_size"),
                page_number=hf.get("page_number", False),
                page_number_pos=hf.get("page_number_position", "center"),
                different_first=hf.get("different_first_page", False),
            )

        # Sections
        sections = spec.get("sections", [])
        for sspec in sections:
            self.add_section_break(
                after_text=sspec.get("break_after_text"),
                after_paragraph=sspec.get("break_after_paragraph"),
                page_number_start=sspec.get("page_number_start"),
            )

        # TOC
        toc = spec.get("toc", {})
        if toc:
            self.insert_toc(
                position=toc.get("position", 0),
                levels=toc.get("levels", 3),
                title=toc.get("title", "目录"),
            )

    # ── report ────────────────────────────────────────────────────────────

    def report(self):
        """Print a summary of all changes."""
        if not self.changes:
            print("\nNo changes applied.")
            return
        print(f"\n{'='*60}")
        print(f"SUMMARY: {len(self.changes)} changes applied to {self.filepath}")
        print(f"{'='*60}")
        for i, c in enumerate(self.changes, 1):
            print(f"  {i:3d}. {c}")
        print(f"\nBackup at: {getattr(self, 'backup_path', 'none')}")
        print(f"Run normalize_fonts.py next if font changes were made.")


# ─── CLI ─────────────────────────────────────────────────────────────────────

def build_parser():
    parser = argparse.ArgumentParser(
        description="Academic paper formatting engine for Word documents.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Page setup
  %(prog)s paper.docx page-setup --size A4 --margins "2.54cm,2.54cm,3.17cm,3.17cm"

  # Modify Normal style (affects all body text)
  %(prog)s paper.docx styles --name Normal --cn-font 宋体 --en-font "Times New Roman" \\
    --font-size 小四 --line-spacing 1.5 --first-indent 2char --alignment justify

  # Modify Heading 1 style
  %(prog)s paper.docx styles --name "Heading 1" --cn-font 黑体 --en-font Arial \\
    --font-size 三号 --bold --space-before 24pt --space-after 12pt --alignment center

  # Set header and footer with page numbers
  %(prog)s paper.docx header-footer --header "My Paper Title" --page-number \\
    --header-font "宋体,Times New Roman" --header-size 五号

  # Apply all formatting from a JSON spec
  %(prog)s paper.docx apply-spec --spec format_spec.json

  # Insert TOC before the paragraph containing "引言"
  %(prog)s paper.docx toc --position "引言" --levels 3 --title 目录
""")

    parser.add_argument("file", help="Path to the .docx file")
    parser.add_argument("--no-backup", action="store_true", help="Skip creating a backup")

    sub = parser.add_subparsers(dest="command", required=True)

    # page-setup
    ps = sub.add_parser("page-setup", help="Set page size, margins, orientation")
    ps.add_argument("--size", help="Paper size: A4, Letter, B5, A3, Legal, 16K")
    ps.add_argument("--margins", help="Margins: 'top,bottom,left,right' in cm/pt/in (e.g. '2.54cm,2.54cm,3.17cm,3.17cm')")
    ps.add_argument("--orientation", help="portrait or landscape")

    # styles
    st = sub.add_parser("styles", help="Modify a built-in or custom style definition")
    st.add_argument("--name", required=True, help="Style name (e.g. Normal, 'Heading 1')")
    st.add_argument("--cn-font", help="Chinese font (e.g. 宋体, 黑体)")
    st.add_argument("--en-font", help="English font (e.g. 'Times New Roman', Arial)")
    st.add_argument("--font-size", help="Font size (e.g. 12pt, 小四)")
    st.add_argument("--bold", action="store_true", default=None, help="Bold")
    st.add_argument("--no-bold", dest="bold", action="store_false", help="Not bold")
    st.add_argument("--italic", action="store_true", default=None, help="Italic")
    st.add_argument("--no-italic", dest="italic", action="store_false", help="Not italic")
    st.add_argument("--color", help="Text color as hex RGB (e.g. 000000)")
    st.add_argument("--alignment", help="Paragraph alignment: left, center, right, justify")
    st.add_argument("--line-spacing", help="Line spacing multiplier (e.g. 1.5, 2.0)")
    st.add_argument("--line-spacing-rule", help="Line spacing rule: multiple (default), exact, at_least")
    st.add_argument("--space-before", help="Space before paragraph (e.g. 12pt, 0.5cm)")
    st.add_argument("--space-after", help="Space after paragraph (e.g. 6pt)")
    st.add_argument("--first-indent", help="First line indent (e.g. 2char, 0.85cm, 24pt)")
    st.add_argument("--hanging-indent", help="Hanging indent (e.g. 1cm)")
    st.add_argument("--keep-with-next", action="store_true", default=None, help="Keep with next paragraph")

    # paragraph
    pg = sub.add_parser("paragraph", help="Apply formatting to paragraphs by filter")
    pg.add_argument("--style", help="Only affect paragraphs with this style name")
    pg.add_argument("--range", help="Paragraph index range: 'start,end' (0-based)")
    pg.add_argument("--alignment", help="Paragraph alignment")
    pg.add_argument("--line-spacing", help="Line spacing multiplier")
    pg.add_argument("--space-before", help="Space before")
    pg.add_argument("--space-after", help="Space after")
    pg.add_argument("--first-indent", help="First line indent")

    # header-footer
    hf = sub.add_parser("header-footer", help="Set header/footer content and page numbers")
    hf.add_argument("--header", dest="header_text", help="Header text")
    hf.add_argument("--footer", dest="footer_text", help="Footer text")
    hf.add_argument("--header-font", help="Header font: 'cn_font,en_font' or single font name")
    hf.add_argument("--header-size", help="Header font size")
    hf.add_argument("--footer-font", help="Footer font")
    hf.add_argument("--footer-size", help="Footer font size")
    hf.add_argument("--page-number", action="store_true", help="Add page number to footer")
    hf.add_argument("--page-number-pos", default="center", help="Page number position: left, center, right")
    hf.add_argument("--different-first", action="store_true", help="Different first page header/footer")

    # section
    sc = sub.add_parser("section", help="Add section break and configure")
    sc.add_argument("--break-after-text", help="Insert section break after paragraph containing this text")
    sc.add_argument("--break-after-paragraph", type=int, help="Insert section break after paragraph index")
    sc.add_argument("--page-number-start", type=int, help="Restart page numbering from this value")

    # toc
    tc = sub.add_parser("toc", help="Insert Table of Contents")
    tc.add_argument("--position", default="0", help="Insert before paragraph index or text match")
    tc.add_argument("--levels", type=int, default=3, help="Heading levels to include (default: 3)")
    tc.add_argument("--title", default="目录", help="TOC title text")

    # apply-spec
    ap = sub.add_parser("apply-spec", help="Apply all formatting from a JSON spec file")
    ap.add_argument("--spec", required=True, help="Path to JSON format spec file")

    return parser


def main():
    parser = build_parser()
    args = parser.parse_args()

    if not os.path.exists(args.file):
        print(f"ERROR: File not found: {args.file}", file=sys.stderr)
        sys.exit(1)

    engine = FormatEngine(args.file, backup=not args.no_backup)

    if args.command == "page-setup":
        engine.page_setup(size=args.size, margins=args.margins, orientation=args.orientation)

    elif args.command == "styles":
        engine.modify_style(
            args.name,
            cn_font=args.cn_font, en_font=args.en_font,
            font_size=args.font_size, bold=args.bold, italic=args.italic,
            color=args.color, alignment=args.alignment,
            line_spacing=args.line_spacing, line_spacing_rule=args.line_spacing_rule,
            space_before=args.space_before, space_after=args.space_after,
            first_indent=args.first_indent, hanging_indent=args.hanging_indent,
            keep_with_next=args.keep_with_next,
        )

    elif args.command == "paragraph":
        idx_range = None
        if args.range:
            parts = args.range.split(",")
            idx_range = (int(parts[0]), int(parts[1]))
        engine.paragraph_format_batch(
            style_filter=args.style, index_range=idx_range,
            alignment=args.alignment, line_spacing=args.line_spacing,
            space_before=args.space_before, space_after=args.space_after,
            first_indent=args.first_indent,
        )

    elif args.command == "header-footer":
        engine.header_footer(
            header_text=args.header_text, footer_text=args.footer_text,
            header_font=args.header_font, header_size=args.header_size,
            footer_font=args.footer_font, footer_size=args.footer_size,
            page_number=args.page_number, page_number_pos=args.page_number_pos,
            different_first=args.different_first,
        )

    elif args.command == "section":
        engine.add_section_break(
            after_text=args.break_after_text,
            after_paragraph=args.break_after_paragraph,
            page_number_start=args.page_number_start,
        )

    elif args.command == "toc":
        pos = args.position
        try:
            pos = int(pos)
        except ValueError:
            pass
        engine.insert_toc(position=pos, levels=args.levels, title=args.title)

    elif args.command == "apply-spec":
        engine.apply_spec(args.spec)

    engine.save()
    engine.report()


if __name__ == "__main__":
    main()
