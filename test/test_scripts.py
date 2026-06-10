"""Regression tests for the script layer (format_document.py, normalize_fonts.py,
accept_changes.py, office/unpack.py + pack.py).

Each test asserts on the resulting XML / reopened document, not just exit codes —
these are the tests that would have caught the transplant-overwrite and
clear-direct-format corruption bugs.
"""
import json
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest

SCRIPTS = Path(__file__).parent.parent / "scripts"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

sys.path.insert(0, str(SCRIPTS))


def run_cli(*argv):
    result = subprocess.run([sys.executable] + [str(a) for a in argv],
                            capture_output=True, text=True)
    return result


def read_part(docx_path, part="word/document.xml"):
    with zipfile.ZipFile(str(docx_path), "r") as z:
        return z.read(part).decode("utf-8")


def doc_text(docx_path):
    from docx import Document
    doc = Document(str(docx_path))
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# transplant-styles (regression for: result overwritten by engine.save())
# ---------------------------------------------------------------------------

class TestTransplantStyles:
    def _make_template(self, sample_docx, output_dir):
        template = output_dir / "template.docx"
        shutil.copy2(sample_docx, template)
        r = run_cli(SCRIPTS / "format_document.py", template, "--no-backup",
                    "styles", "--name", "Normal", "--font-size", "13pt",
                    "--cn-font", "楷体", "--en-font", "Georgia")
        assert r.returncode == 0, r.stderr
        return template

    def test_transplant_persists_on_disk(self, sample_docx, output_dir):
        """The transplanted styles.xml must survive — not be clobbered by a
        stale in-memory save."""
        template = self._make_template(sample_docx, output_dir)
        target = output_dir / "target.docx"
        shutil.copy2(sample_docx, target)

        r = run_cli(SCRIPTS / "format_document.py", target, "--no-backup",
                    "transplant-styles", "--template", template)
        assert r.returncode == 0, r.stderr

        styles = read_part(target, "word/styles.xml")
        # 13pt → sz val 26 half-points; fonts from the template style
        assert 'w:val="26"' in styles
        assert "楷体" in styles
        assert "Georgia" in styles

    def test_transplant_output_reopens(self, sample_docx, output_dir):
        template = self._make_template(sample_docx, output_dir)
        target = output_dir / "target2.docx"
        shutil.copy2(sample_docx, target)
        run_cli(SCRIPTS / "format_document.py", target, "--no-backup",
                "transplant-styles", "--template", template, "--page-setup")
        from docx import Document
        doc = Document(str(target))  # raises if corrupt
        assert doc.paragraphs

    def test_transplant_page_setup_keeps_declaration(self, sample_docx, output_dir):
        """--page-setup must not strip the XML declaration from document.xml."""
        template = self._make_template(sample_docx, output_dir)
        target = output_dir / "target3.docx"
        shutil.copy2(sample_docx, target)
        run_cli(SCRIPTS / "format_document.py", target, "--no-backup",
                "transplant-styles", "--template", template, "--page-setup")
        xml = read_part(target)
        assert xml.lstrip().startswith("<?xml")
        # header/footer references from the template's rels must not be copied
        body_sect = xml.rsplit("<w:sectPr", 1)[-1]
        assert "headerReference" not in body_sect
        assert "footerReference" not in body_sect


# ---------------------------------------------------------------------------
# clear-direct-format (regression for: sectPr/revision/rStyle destruction)
# ---------------------------------------------------------------------------

class TestClearDirectFormat:
    def _clear(self, docx):
        r = run_cli(SCRIPTS / "format_document.py", docx, "--no-backup",
                    "clear-direct-format")
        assert r.returncode == 0, r.stderr

    def test_strips_direct_formatting(self, sample_docx):
        before = read_part(sample_docx)
        assert "<w:b/>" in before  # the bold direct run exists
        self._clear(sample_docx)
        after = read_part(sample_docx)
        # Direct bold and paragraph-level centering are gone…
        # (jc may remain inside ins/del or styles part; check the cleared para)
        assert "<w:b/>" not in after

    def test_preserves_sections(self, sample_docx):
        from docx import Document
        n_before = len(Document(str(sample_docx)).sections)
        assert n_before == 2
        self._clear(sample_docx)
        assert len(Document(str(sample_docx)).sections) == n_before

    def test_preserves_superscript_and_revisions(self, sample_docx):
        before = read_part(sample_docx)
        assert "vertAlign" in before
        assert "Georgia" in before  # font inside w:ins
        self._clear(sample_docx)
        after = read_part(sample_docx)
        assert "vertAlign" in after, "superscript (m²) was destroyed"
        assert "<w:ins " in after and "<w:del " in after, "tracked changes lost"
        assert "Georgia" in after, "run formatting inside tracked change was altered"

    def test_output_reopens(self, sample_docx):
        self._clear(sample_docx)
        from docx import Document
        assert Document(str(sample_docx)).paragraphs


# ---------------------------------------------------------------------------
# normalize_fonts.py
# ---------------------------------------------------------------------------

class TestNormalizeFonts:
    def test_inplace_creates_backup(self, sample_docx, output_dir):
        r = run_cli(SCRIPTS / "normalize_fonts.py", sample_docx)
        assert r.returncode in (0, 3), r.stderr
        backups = list(output_dir.glob("test_copy_backup_*.docx"))
        assert backups, "in-place normalization must create a backup"

    def test_output_reopens_and_pairs_fonts(self, sample_docx):
        r = run_cli(SCRIPTS / "normalize_fonts.py", sample_docx, "--no-backup", "--json")
        assert r.returncode in (0, 3), r.stderr
        data = json.loads(r.stdout)
        assert data["runs_fixed"] >= 0
        from docx import Document
        assert Document(str(sample_docx)).paragraphs

    def test_unify_preserves_symbol_fonts(self, sample_docx):
        r = run_cli(SCRIPTS / "normalize_fonts.py", sample_docx, "--no-backup", "--unify")
        assert r.returncode in (0, 3), r.stderr
        after = read_part(sample_docx)
        assert "Wingdings" in after, "--unify must not flatten symbol fonts"

    def test_skip_revisions_leaves_tracked_runs(self, sample_docx):
        r = run_cli(SCRIPTS / "normalize_fonts.py", sample_docx, "--no-backup",
                    "--unify", "--skip-revisions")
        assert r.returncode in (0, 3), r.stderr
        after = read_part(sample_docx)
        assert 'w:ascii="Georgia"' in after, \
            "--skip-revisions must leave runs inside w:ins untouched"

    def test_unify_rewrites_revision_runs_without_flag(self, sample_docx):
        r = run_cli(SCRIPTS / "normalize_fonts.py", sample_docx, "--no-backup", "--unify")
        assert r.returncode in (0, 3), r.stderr
        after = read_part(sample_docx)
        assert 'w:ascii="Georgia"' not in after

    def test_theme_ea_fix_is_surgical(self):
        """Only <a:ea> slots change; <a:latin> and compound names survive."""
        from normalize_fonts import _fix_theme_ea_fonts
        xml = (b'<a:latin typeface="Calibri Light"/>'
               b'<a:ea typeface="DengXian"/>'
               b'<a:ea typeface="DejaVu Sans"/>'
               b'<a:ea typeface=""/>'
               b'<a:ea typeface="\xe5\xae\x8b\xe4\xbd\x93"/>')  # 宋体
        out, n = _fix_theme_ea_fonts(xml, "宋体")
        assert n == 1  # only DejaVu Sans is non-CJK
        assert b'<a:latin typeface="Calibri Light"/>' in out, "latin slot must be untouched"
        assert b"DejaVu Sans" not in out
        # DengXian (等线) IS a CJK font — a legitimate ea value, left alone
        assert b'<a:ea typeface="DengXian"/>' in out
        assert b'<a:ea typeface=""/>' in out, "empty ea slot must be left alone"


# ---------------------------------------------------------------------------
# accept_changes.py
# ---------------------------------------------------------------------------

class TestAcceptChanges:
    def test_accepts_ins_and_del(self, sample_docx, output_dir):
        out = output_dir / "clean.docx"
        r = run_cli(SCRIPTS / "accept_changes.py", sample_docx, out)
        assert r.returncode == 0, r.stderr
        xml = read_part(out)
        assert "<w:ins " not in xml and "<w:del " not in xml
        text = doc_text(out)
        assert "inserted text" in text
        assert "deleted text" not in text

    def test_output_reopens(self, sample_docx, output_dir):
        out = output_dir / "clean2.docx"
        run_cli(SCRIPTS / "accept_changes.py", sample_docx, out)
        from docx import Document
        assert Document(str(out)).paragraphs


# ---------------------------------------------------------------------------
# office/unpack.py + pack.py roundtrip
# ---------------------------------------------------------------------------

class TestUnpackPackRoundtrip:
    def test_roundtrip_preserves_document(self, sample_docx, output_dir):
        workdir = output_dir / "unpacked"
        if workdir.exists():
            shutil.rmtree(workdir)
        r = run_cli(SCRIPTS / "office" / "unpack.py", sample_docx, workdir)
        assert r.returncode == 0, r.stderr
        assert (workdir / "word" / "document.xml").exists()

        repacked = output_dir / "repacked.docx"
        r = run_cli(SCRIPTS / "office" / "pack.py", workdir, repacked)
        assert r.returncode == 0, r.stderr

        assert doc_text(repacked) == doc_text(sample_docx)


# ---------------------------------------------------------------------------
# format_document.py misc regressions
# ---------------------------------------------------------------------------

class TestFormatDocumentMisc:
    def test_en_font_auto_pairs_chinese(self, sample_docx):
        """--en-font alone must auto-set a Chinese eastAsia partner (the
        reversed-table-lookup bug left eastAsia unset)."""
        r = run_cli(SCRIPTS / "format_document.py", sample_docx, "--no-backup",
                    "styles", "--name", "Normal", "--en-font", "Times New Roman")
        assert r.returncode == 0, r.stderr
        styles = read_part(sample_docx, "word/styles.xml")
        normal = styles.split('w:styleId="Normal"', 1)[1].split("</w:style>", 1)[0]
        assert 'w:eastAsia="宋体"' in normal

    def test_split_font_never_puts_latin_in_eastasia(self):
        from format_document import FormatEngine
        cn, en = FormatEngine._split_font(FormatEngine, "Arial")
        assert cn == "黑体" and en == "Arial"
        cn, en = FormatEngine._split_font(FormatEngine, "宋体")
        assert cn == "宋体" and en == "Times New Roman"
