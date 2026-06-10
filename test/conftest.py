"""Shared test fixtures for word-agent test suite."""
import os
import shutil
import pytest
from pathlib import Path

FIXTURES_DIR = Path(__file__).parent / "fixtures"
TEST_OUTPUT_DIR = Path(__file__).parent / "output"


@pytest.fixture(autouse=True)
def setup_output_dir():
    """Create and clean output directory for each test."""
    TEST_OUTPUT_DIR.mkdir(exist_ok=True)
    yield
    # Keep output for inspection; clean on next run


@pytest.fixture
def fixtures_dir():
    return FIXTURES_DIR


@pytest.fixture
def output_dir():
    return TEST_OUTPUT_DIR


def _build_sample_paper(path):
    """Generate the sample fixture programmatically: a bilingual paper with
    two sections, direct formatting, a superscript, a symbol font run, and
    injected tracked changes (w:ins with Georgia font + w:del)."""
    import re
    import zipfile
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("测试论文 Test Paper", 0)
    doc.add_heading("1 引言 Introduction", 1)
    p = doc.add_paragraph("这是中文段落，混合 English text 和数字 123。")
    r = p.add_run("直接格式化的 run")
    r.bold = True
    r.font.name = "Arial"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # direct paragraph formatting

    doc.add_paragraph("SENTINEL_TRACKED_HERE")

    p3 = doc.add_paragraph("面积单位 m")
    sup = p3.add_run("2")
    sup.font.superscript = True

    sym = doc.add_paragraph().add_run("")
    sym.font.name = "Wingdings"

    doc.add_heading("2 方法 Methods", 1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头"
    table.cell(1, 1).text = "数据 data"

    sec = doc.add_section()
    sec.orientation = WD_ORIENT.LANDSCAPE
    doc.add_paragraph("第二节内容 second section text")
    doc.save(str(path))

    # Inject tracked changes by replacing the sentinel run at XML level.
    with zipfile.ZipFile(str(path), "r") as z:
        entries = {n: z.read(n) for n in z.namelist()}
    docxml = entries["word/document.xml"].decode("utf-8")
    tracked = (
        '<w:ins w:id="901" w:author="Tester" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:rPr><w:rFonts w:ascii="Georgia" w:hAnsi="Georgia"/></w:rPr>'
        "<w:t>inserted text</w:t></w:r></w:ins>"
        '<w:del w:id="902" w:author="Tester" w:date="2026-01-01T00:00:00Z">'
        '<w:r><w:delText xml:space="preserve">deleted text</w:delText></w:r></w:del>'
    )
    new_docxml, n = re.subn(
        r"<w:r>(?:(?!</w:r>).)*?SENTINEL_TRACKED_HERE.*?</w:r>",
        tracked, docxml, count=1, flags=re.DOTALL)
    assert n == 1, "sentinel run not found while building fixture"
    entries["word/document.xml"] = new_docxml.encode("utf-8")
    with zipfile.ZipFile(str(path), "w") as z:
        for name, data in entries.items():
            z.writestr(name, data)


@pytest.fixture
def sample_docx(fixtures_dir, output_dir):
    """Copy a sample docx to output dir for modification tests.
    The fixture file is generated programmatically on first use."""
    src = fixtures_dir / "sample_paper.docx"
    if not src.exists():
        try:
            _build_sample_paper(src)
        except ImportError:
            pytest.skip("python-docx not installed; cannot build fixture")
    dst = output_dir / "test_copy.docx"
    shutil.copy2(src, dst)
    return dst


@pytest.fixture
def scripts_dir():
    return Path(__file__).parent.parent / "scripts"


@pytest.fixture
def normalize_fonts_script(scripts_dir):
    script = scripts_dir / "normalize_fonts.py"
    if not script.exists():
        pytest.skip("normalize_fonts.py not found")
    return script
